import io
import shutil
import zipfile

import pytest
from PIL import Image
from pypdf import PdfReader, PdfWriter
from werkzeug.datastructures import FileStorage

from app_factory import create_app


def upload(name, content):
    return (io.BytesIO(content), name)


def make_pdf(pages=1):
    stream = io.BytesIO()
    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=144, height=144)
    writer.add_metadata({"/Title": "Infinity test"})
    writer.write(stream)
    stream.seek(0)
    return stream


def make_image(format_name, mode="RGBA"):
    stream = io.BytesIO()
    image = Image.new(mode, (24, 16), (20, 120, 70, 128) if mode == "RGBA" else "white")
    image.save(stream, format=format_name)
    stream.seek(0)
    return stream


def test_pdf_engines_validate_real_outputs():
    client = create_app().test_client()
    response = client.post(
        "/api/v2/convert",
        data={"tool": "pdf-merge", "files": [upload("one.pdf", make_pdf().read()), upload("two.pdf", make_pdf(2).read())]},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    assert len(PdfReader(io.BytesIO(response.data)).pages) == 3

    response = client.post(
        "/api/v2/convert",
        data={"tool": "pdf-split", "file": upload("source.pdf", make_pdf(2).read())},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    with zipfile.ZipFile(io.BytesIO(response.data)) as zf:
        names = zf.namelist()
        assert len(names) == 2
        assert len(PdfReader(io.BytesIO(zf.read(names[0]))).pages) == 1


def test_image_engines_validate_real_outputs():
    client = create_app().test_client()
    response = client.post(
        "/api/v2/convert",
        data={"tool": "image-to-jpg", "file": upload("source.png", make_image("PNG").read())},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    with Image.open(io.BytesIO(response.data)) as image:
        assert image.format == "JPEG"
        assert image.mode == "RGB"

    response = client.post(
        "/api/v2/convert",
        data={"tool": "image-to-png", "file": upload("source.jpg", make_image("JPEG", "RGB").read())},
        content_type="multipart/form-data",
    )
    assert response.status_code == 200
    with Image.open(io.BytesIO(response.data)) as image:
        assert image.format == "PNG"


def test_tool_rejects_valid_but_wrong_extension():
    client = create_app().test_client()
    response = client.post(
        "/api/v2/convert",
        data={"tool": "image-to-jpg", "file": upload("source.pdf", make_pdf().read())},
        content_type="multipart/form-data",
    )
    assert response.status_code == 400
    assert "تقبل" in response.get_json()["error"]


@pytest.mark.skipif(shutil.which("libreoffice") is None, reason="LibreOffice is required for Office integration tests")
def test_office_engine_converts_docx(tmp_path):
    from docx import Document

    source = tmp_path / "source.docx"
    document = Document()
    document.add_heading("Infinity conversion", level=1)
    document.add_paragraph("Arabic content: اختبار التحويل")
    document.save(source)

    client = create_app().test_client()
    with source.open("rb") as stream:
        response = client.post(
            "/api/v2/convert",
            data={"tool": "word-to-pdf", "file": FileStorage(stream=stream, filename="source.docx")},
            content_type="multipart/form-data",
        )
    assert response.status_code == 200
    assert len(PdfReader(io.BytesIO(response.data)).pages) >= 1
